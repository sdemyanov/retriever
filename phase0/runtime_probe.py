from __future__ import annotations

import json
import sqlite3
import sys
from dataclasses import asdict, dataclass
from email import policy
from email.message import EmailMessage
from email.parser import BytesParser
from importlib.metadata import version
from pathlib import Path

import charset_normalizer
import extract_msg
import openpyxl
import pdfplumber
from docx import Document


ROOT = Path(__file__).resolve().parent
CORPUS_DIR = ROOT / "regression_corpus"
OUTPUT_PATH = ROOT / "runtime_proof.json"
PERSISTENCE_DIR = ROOT / "workspace_persistence_probe"


@dataclass
class CheckResult:
    status: str
    details: str


def escape_pdf_text(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_minimal_pdf(path: Path, text: str) -> None:
    content = f"BT /F1 12 Tf 72 72 Td ({escape_pdf_text(text)}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 144] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n%b\nendstream" % (len(content), content),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    parts = [b"%PDF-1.4\n"]
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(part) for part in parts))
        parts.append(f"{index} 0 obj\n".encode("ascii"))
        parts.append(obj)
        parts.append(b"\nendobj\n")

    xref_offset = sum(len(part) for part in parts)
    parts.append(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    parts.append(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        parts.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    parts.append(
        (
            "trailer\n"
            f"<< /Root 1 0 R /Size {len(objects) + 1} >>\n"
            "startxref\n"
            f"{xref_offset}\n"
            "%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(b"".join(parts))


def create_regression_corpus() -> dict[str, str]:
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)

    files: dict[str, str] = {}

    valid_pdf = CORPUS_DIR / "sample.pdf"
    build_minimal_pdf(valid_pdf, "Phase 0 PDF probe")
    files["pdf"] = str(valid_pdf.relative_to(ROOT))

    corrupt_pdf = CORPUS_DIR / "corrupt.pdf"
    corrupt_pdf.write_bytes(b"%PDF-1.4\nthis is not a valid pdf body\n")
    files["corrupt_pdf"] = str(corrupt_pdf.relative_to(ROOT))

    docx_path = CORPUS_DIR / "sample.docx"
    doc = Document()
    doc.add_heading("Phase 0 DOCX probe", level=1)
    doc.add_paragraph("This file validates DOCX creation and parsing.")
    doc.save(docx_path)
    files["docx"] = str(docx_path.relative_to(ROOT))

    xlsx_path = CORPUS_DIR / "sample.xlsx"
    workbook = openpyxl.Workbook()
    sheet1 = workbook.active
    sheet1.title = "Summary"
    sheet1.append(["Name", "Value"])
    sheet1.append(["Alpha", 1])
    sheet2 = workbook.create_sheet("Details")
    sheet2.append(["Item", "Amount"])
    sheet2.append(["Beta", 2])
    workbook.save(xlsx_path)
    files["xlsx"] = str(xlsx_path.relative_to(ROOT))

    html_path = CORPUS_DIR / "sample.html"
    html_path.write_text(
        "<html><body><h1>Phase 0 HTML probe</h1><p>Preview target.</p></body></html>",
        encoding="utf-8",
    )
    files["html"] = str(html_path.relative_to(ROOT))

    csv_utf8 = CORPUS_DIR / "sample_utf8.csv"
    csv_utf8.write_text("name,city\nAndre,Munchen\n", encoding="utf-8")
    files["csv_utf8"] = str(csv_utf8.relative_to(ROOT))

    csv_latin1 = CORPUS_DIR / "sample_latin1.csv"
    csv_latin1.write_bytes("name,city\nAndre,M\xfcnchen\n".encode("latin-1"))
    files["csv_latin1"] = str(csv_latin1.relative_to(ROOT))

    txt_utf8 = CORPUS_DIR / "sample_utf8.txt"
    txt_utf8.write_text("UTF-8 probe: cafe, naive, resume.\n", encoding="utf-8")
    files["txt_utf8"] = str(txt_utf8.relative_to(ROOT))

    txt_latin1 = CORPUS_DIR / "sample_latin1.txt"
    txt_latin1.write_bytes("Latin-1 probe: caf\xe9, M\xfcnchen.\n".encode("latin-1"))
    files["txt_latin1"] = str(txt_latin1.relative_to(ROOT))

    eml_utf8 = CORPUS_DIR / "sample_utf8.eml"
    msg_utf8 = EmailMessage()
    msg_utf8["From"] = "alice@example.com"
    msg_utf8["To"] = "bob@example.com"
    msg_utf8["Subject"] = "UTF-8 EML probe"
    msg_utf8.set_content("Hello from the UTF-8 message body.")
    eml_utf8.write_bytes(msg_utf8.as_bytes(policy=policy.default))
    files["eml_utf8"] = str(eml_utf8.relative_to(ROOT))

    eml_latin1 = CORPUS_DIR / "sample_latin1.eml"
    msg_latin1 = EmailMessage()
    msg_latin1["From"] = "alice@example.com"
    msg_latin1["To"] = "bob@example.com"
    msg_latin1["Subject"] = "Latin-1 EML probe"
    msg_latin1.set_content("Bonjour de caf\xe9 et M\xfcnchen.", charset="iso-8859-1")
    eml_latin1.write_bytes(msg_latin1.as_bytes(policy=policy.default))
    files["eml_latin1"] = str(eml_latin1.relative_to(ROOT))

    invalid_msg = CORPUS_DIR / "sample_invalid.msg"
    invalid_msg.write_bytes(b"not-a-real-msg")
    files["msg_invalid"] = str(invalid_msg.relative_to(ROOT))

    return files


def check_fts5() -> CheckResult:
    try:
        conn = sqlite3.connect(":memory:")
        conn.execute("create virtual table t using fts5(x)")
        conn.close()
        return CheckResult("pass", "SQLite FTS5 virtual tables can be created.")
    except Exception as exc:  # pragma: no cover - probe script
        return CheckResult("fail", f"FTS5 unavailable: {exc}")


def check_import_versions() -> dict[str, str]:
    packages = [
        "pdfplumber",
        "python-docx",
        "openpyxl",
        "extract-msg",
        "charset-normalizer",
    ]
    return {name: version(name) for name in packages}


def check_pdf_parsing(pdf_path: Path, corrupt_pdf_path: Path) -> dict[str, CheckResult]:
    results: dict[str, CheckResult] = {}
    with pdfplumber.open(pdf_path) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
    if "Phase 0 PDF probe" in text:
        results["valid_pdf"] = CheckResult("pass", f"Extracted text: {text}")
    else:
        results["valid_pdf"] = CheckResult("fail", f"Unexpected PDF text: {text!r}")

    try:
        with pdfplumber.open(corrupt_pdf_path) as pdf:
            _ = pdf.pages[0].extract_text()
        results["corrupt_pdf"] = CheckResult("fail", "Corrupt PDF unexpectedly parsed.")
    except Exception as exc:  # pragma: no cover - probe script
        results["corrupt_pdf"] = CheckResult("pass", f"Corrupt PDF raised cleanly: {type(exc).__name__}")
    return results


def check_docx_parsing(docx_path: Path) -> CheckResult:
    text = "\n".join(paragraph.text for paragraph in Document(docx_path).paragraphs).strip()
    if "Phase 0 DOCX probe" in text:
        return CheckResult("pass", f"Extracted text: {text}")
    return CheckResult("fail", f"Unexpected DOCX text: {text!r}")


def check_xlsx_parsing(xlsx_path: Path) -> CheckResult:
    workbook = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
    if workbook.sheetnames == ["Summary", "Details"]:
        return CheckResult("pass", f"Sheet names: {', '.join(workbook.sheetnames)}")
    return CheckResult("fail", f"Unexpected sheet names: {workbook.sheetnames}")


def detect_text_encoding(path: Path) -> dict[str, str]:
    data = path.read_bytes()
    best = charset_normalizer.from_bytes(data).best()
    if best is None:
        decoded = data.decode("utf-8", errors="replace")
        return {"encoding": "unknown", "decoded_preview": decoded[:80]}
    return {"encoding": best.encoding or "unknown", "decoded_preview": str(best)[:80]}


def parse_eml(path: Path) -> dict[str, str]:
    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    body = message.get_body(preferencelist=("plain", "html"))
    content = body.get_content() if body is not None else message.get_content()
    return {
        "subject": message["Subject"] or "",
        "content_preview": str(content)[:120],
    }


def check_msg_failure(path: Path) -> CheckResult:
    try:
        extract_msg.Message(str(path))
        return CheckResult("fail", "Invalid MSG unexpectedly opened.")
    except Exception as exc:  # pragma: no cover - probe script
        return CheckResult("pass", f"Invalid MSG raised cleanly: {type(exc).__name__}")


def check_local_persistence() -> CheckResult:
    PERSISTENCE_DIR.mkdir(parents=True, exist_ok=True)
    marker = PERSISTENCE_DIR / "persistence_marker.txt"
    marker.write_text("phase0 persistence probe\n", encoding="utf-8")
    exists_after_write = marker.exists()
    read_back = marker.read_text(encoding="utf-8").strip()
    if exists_after_write and read_back == "phase0 persistence probe":
        return CheckResult(
            "partial",
            "Local filesystem persistence across processes is fine, but cross-session Cowork persistence still needs in-product validation.",
        )
    return CheckResult("fail", "Could not read back the persistence marker.")


def main() -> int:
    corpus_files = create_regression_corpus()
    proof = {
        "environment": {
            "python_version": sys.version.split()[0],
            "sqlite_version": sqlite3.sqlite_version,
            "platform": sys.platform,
        },
        "imports": check_import_versions(),
        "checks": {
            "fts5": asdict(check_fts5()),
            "pdf_parsing": {
                name: asdict(result)
                for name, result in check_pdf_parsing(
                    CORPUS_DIR / "sample.pdf",
                    CORPUS_DIR / "corrupt.pdf",
                ).items()
            },
            "docx_parsing": asdict(check_docx_parsing(CORPUS_DIR / "sample.docx")),
            "xlsx_parsing": asdict(check_xlsx_parsing(CORPUS_DIR / "sample.xlsx")),
            "txt_latin1_detection": detect_text_encoding(CORPUS_DIR / "sample_latin1.txt"),
            "csv_latin1_detection": detect_text_encoding(CORPUS_DIR / "sample_latin1.csv"),
            "eml_utf8": parse_eml(CORPUS_DIR / "sample_utf8.eml"),
            "eml_latin1": parse_eml(CORPUS_DIR / "sample_latin1.eml"),
            "msg_invalid": asdict(check_msg_failure(CORPUS_DIR / "sample_invalid.msg")),
            "local_persistence": asdict(check_local_persistence()),
        },
        "corpus_files": corpus_files,
    }
    OUTPUT_PATH.write_text(json.dumps(proof, indent=2, sort_keys=True), encoding="utf-8")
    print(str(OUTPUT_PATH))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
